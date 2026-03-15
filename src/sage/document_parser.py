"""Parse uploaded document files (.docx / .txt) into plain text."""

from __future__ import annotations

import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


def parse_file(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from an uploaded file.

    Supported formats:
    - .docx  (via python-docx)
    - .txt   (plain text, auto-detect encoding)
    """
    lower = filename.lower()

    if lower.endswith(".docx"):
        return _parse_docx(file_bytes)
    elif lower.endswith(".txt"):
        return _parse_txt(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {filename}。请上传 .docx 或 .txt 文件。")


def _parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]

    # Also extract text from tables (rubric files often use tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _parse_txt(data: bytes) -> str:
    # Try UTF-8 first, fallback to GBK (common for Chinese text)
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")
