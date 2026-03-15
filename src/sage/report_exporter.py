"""Export grading reports to Word (.docx) and PDF formats.

Handles Chinese fonts and proper Markdown rendering.
"""

from __future__ import annotations

import io
import logging
import re

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

# ── Markdown → HTML (shared) ────────────────────────

_MD_EXTENSIONS = ["tables", "fenced_code", "nl2br"]


def _md_to_html(md_text: str) -> str:
    """Convert Markdown to HTML fragment."""
    return markdown.markdown(md_text, extensions=_MD_EXTENSIONS)


# ── PDF Export ──────────────────────────────────────

_PDF_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@300;400;500;700&display=swap');

* {{ margin: 0; padding: 0; box-sizing: border-box; }}

body {{
    font-family: 'Microsoft YaHei', '微软雅黑', 'Noto Sans SC', 'PingFang SC',
                 'Hiragino Sans GB', 'SimSun', sans-serif;
    font-size: 11pt;
    line-height: 1.8;
    color: #1a1a1a;
    padding: 40px 50px;
}}

h1 {{
    font-size: 18pt;
    color: #2c3e50;
    border-bottom: 2px solid #3498db;
    padding-bottom: 8px;
    margin-bottom: 16px;
}}

h2 {{
    font-size: 14pt;
    color: #2c3e50;
    margin-top: 22px;
    margin-bottom: 8px;
}}

h3 {{
    font-size: 12pt;
    color: #3498db;
    margin-top: 16px;
    margin-bottom: 6px;
}}

p {{
    margin: 6px 0;
    text-align: justify;
}}

ul, ol {{
    margin: 6px 0 6px 24px;
}}

li {{
    margin: 3px 0;
}}

strong {{
    color: #2c3e50;
}}

blockquote {{
    border-left: 3px solid #3498db;
    padding-left: 14px;
    margin: 10px 0;
    color: #555;
    font-style: italic;
}}

code {{
    background: #f0f0f0;
    padding: 2px 5px;
    border-radius: 3px;
    font-size: 10pt;
}}

hr {{
    border: none;
    border-top: 1px solid #ddd;
    margin: 18px 0;
}}

table {{
    width: 100%;
    border-collapse: collapse;
    margin: 10px 0;
}}

th, td {{
    border: 1px solid #ddd;
    padding: 6px 10px;
    text-align: left;
}}

th {{
    background: #f5f5f5;
    font-weight: 600;
}}
</style>
</head>
<body>
{body}
</body>
</html>
"""


def export_pdf(md_text: str, title: str = "", author: str = "") -> bytes:
    """Convert Markdown report to PDF bytes with Chinese font support."""
    from weasyprint import HTML

    html_body = _md_to_html(md_text)

    # Add a header if title/author provided but not in the markdown
    if title and not md_text.strip().startswith("#"):
        header = f"<h1>{title}</h1>"
        if author:
            header += f"<p><strong>作者：</strong>{author}</p><hr>"
        html_body = header + html_body

    full_html = _PDF_HTML_TEMPLATE.format(body=html_body)

    pdf_bytes = HTML(string=full_html).write_pdf()
    logger.info("Generated PDF: %d bytes", len(pdf_bytes))
    return pdf_bytes


# ── DOCX Export ─────────────────────────────────────


def export_docx(md_text: str, title: str = "", author: str = "") -> bytes:
    """Convert Markdown report to DOCX bytes with Chinese font support."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(11)

    # Set heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Microsoft YaHei"
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Parse and render Markdown line by line
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line → skip
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("###"):
            _add_heading(doc, stripped.lstrip("#").strip(), level=3)
        elif stripped.startswith("##"):
            _add_heading(doc, stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("#"):
            _add_heading(doc, stripped.lstrip("#").strip(), level=1)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
        # Unordered list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        # Ordered list
        elif re.match(r"^\d+[.、]\s", stripped):
            text = re.sub(r"^\d+[.、]\s*", "", stripped)
            _add_rich_paragraph(doc, text, style="List Number")
        # Blockquote
        elif stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            _add_runs(para, text, italic=True)
        # Regular paragraph
        else:
            _add_rich_paragraph(doc, stripped)

        i += 1

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = buf.getvalue()
    logger.info("Generated DOCX: %d bytes", len(result))
    return result


def _add_heading(doc: Document, text: str, level: int):
    """Add a heading with bold/emphasis support."""
    heading = doc.add_heading(level=level)
    _add_runs(heading, text)


def _add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    """Add a paragraph with inline Markdown formatting (bold, italic, code)."""
    para = doc.add_paragraph(style=style)
    _add_runs(para, text)


def _add_runs(paragraph, text: str, italic: bool = False):
    """Parse inline Markdown (bold, italic, code, bold-italic) and add runs."""
    # Pattern: ***bold italic*** or **bold** or *italic* or `code`
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # bold italic
        r"|(\*\*(.+?)\*\*)"       # bold
        r"|(\*(.+?)\*)"           # italic
        r"|(`(.+?)`)"             # inline code
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.italic = italic

        if match.group(2):  # bold italic
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # bold
            run = paragraph.add_run(match.group(4))
            run.bold = True
            run.italic = italic
        elif match.group(6):  # italic
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # code
            run = paragraph.add_run(match.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x40, 0x40)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.italic = italic
